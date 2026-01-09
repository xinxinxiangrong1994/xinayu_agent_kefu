"""生成完整功能方案文档 DOCX"""
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def create_document():
    doc = Document()

    # 设置中文字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 标题
    title = doc.add_heading('闲鱼智能客服 RPA - 系统技术方案文档', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 文档信息
    doc.add_paragraph('版本: v1.1')
    doc.add_paragraph('更新日期: 2026-01-06')
    doc.add_paragraph('文档目的: 记录系统完整技术方案和功能实现细节')
    doc.add_paragraph()

    # ==================== 一、系统概述 ====================
    doc.add_heading('一、系统概述', 1)

    doc.add_heading('1.1 系统简介', 2)
    doc.add_paragraph('闲鱼智能客服 RPA 是一款基于 Playwright 浏览器自动化和 Coze AI 的智能客服系统，用于自动监控和回复闲鱼平台的买家消息。')

    doc.add_heading('1.2 核心能力', 2)
    doc.add_paragraph('1. 自动监控闲鱼网页版消息')
    doc.add_paragraph('2. 智能识别买家消息并调用 AI 生成回复')
    doc.add_paragraph('3. 支持多轮对话记忆（基于 MySQL 存储）')
    doc.add_paragraph('4. 传递订单状态、商品信息等上下文给 AI')
    doc.add_paragraph('5. 可视化 GUI 界面，支持配置管理')
    doc.add_paragraph()

    # ==================== 二、系统架构 ====================
    doc.add_heading('二、系统架构', 1)

    doc.add_heading('2.1 技术栈', 2)

    table_tech = doc.add_table(rows=7, cols=2)
    table_tech.style = 'Table Grid'

    tech_header = table_tech.rows[0].cells
    tech_header[0].text = '技术组件'
    tech_header[1].text = '说明'

    tech_data = [
        ('Python 3.x', '主要开发语言'),
        ('Playwright', '浏览器自动化框架，用于操作闲鱼网页'),
        ('Coze API v3', '扣子 AI 平台，提供智能对话能力'),
        ('MySQL', '数据库，存储用户信息和对话历史'),
        ('Tkinter', 'GUI 界面框架'),
        ('Loguru', '日志系统'),
    ]

    for i, (tech, desc) in enumerate(tech_data, 1):
        row = table_tech.rows[i].cells
        row[0].text = tech
        row[1].text = desc

    doc.add_paragraph()

    doc.add_heading('2.2 模块结构', 2)

    table_module = doc.add_table(rows=9, cols=2)
    table_module.style = 'Table Grid'

    module_header = table_module.rows[0].cells
    module_header[0].text = '文件'
    module_header[1].text = '功能说明'

    module_data = [
        ('gui.py', '可视化界面，配置管理和启动控制'),
        ('xianyu_browser.py', '浏览器自动化，页面操作和数据提取'),
        ('message_handler.py', '消息处理核心逻辑，协调各模块'),
        ('coze_client.py', 'Coze API 客户端，负责 AI 对话'),
        ('db_manager.py', '数据库管理，用户和对话存储'),
        ('config.py', '配置管理，环境变量和常量定义'),
        ('logger_setup.py', '日志配置，文件和控制台输出'),
        ('main.py', '命令行入口'),
    ]

    for i, (file, desc) in enumerate(module_data, 1):
        row = table_module.rows[i].cells
        row[0].text = file
        row[1].text = desc

    doc.add_paragraph()

    # ==================== 三、浏览器自动化 ====================
    doc.add_heading('三、浏览器自动化模块', 1)

    doc.add_heading('3.1 功能概述', 2)
    doc.add_paragraph('使用 Playwright 框架操作闲鱼网页版（goofish.com），实现以下功能：')
    doc.add_paragraph('1. 持久化登录状态（浏览器数据目录）')
    doc.add_paragraph('2. 监控会话列表，识别未读消息')
    doc.add_paragraph('3. 进入会话，提取消息内容')
    doc.add_paragraph('4. 提取商品信息和订单状态')
    doc.add_paragraph('5. 发送回复消息')

    doc.add_heading('3.2 核心类', 2)
    doc.add_paragraph('XianyuBrowser 类提供以下方法：')

    table_browser = doc.add_table(rows=9, cols=2)
    table_browser.style = 'Table Grid'

    browser_header = table_browser.rows[0].cells
    browser_header[0].text = '方法'
    browser_header[1].text = '功能'

    browser_data = [
        ('start()', '启动浏览器，加载持久化上下文'),
        ('navigate_to_messages()', '导航到消息页面'),
        ('check_login_status()', '检查登录状态'),
        ('wait_for_login()', '等待用户手动登录'),
        ('get_conversation_list()', '获取会话列表'),
        ('get_unread_conversations()', '获取未读会话'),
        ('get_current_conversation_messages()', '获取当前会话消息'),
        ('send_message()', '发送消息'),
    ]

    for i, (method, desc) in enumerate(browser_data, 1):
        row = table_browser.rows[i].cells
        row[0].text = method
        row[1].text = desc

    doc.add_paragraph()

    # ==================== 四、Coze AI 集成 ====================
    doc.add_heading('四、Coze AI 集成', 1)

    doc.add_heading('4.1 API 版本', 2)
    doc.add_paragraph('使用 Coze API v3（国内版 api.coze.cn）')

    doc.add_heading('4.2 核心功能', 2)
    doc.add_paragraph('1. 创建会话 - 为每个买家创建独立会话')
    doc.add_paragraph('2. 发送消息 - 调用 Bot 获取 AI 回复')
    doc.add_paragraph('3. 轮询结果 - 异步获取对话结果')
    doc.add_paragraph('4. 传递变量 - 通过 custom_variables 和 parameters 传递上下文')

    doc.add_heading('4.3 会话管理', 2)
    doc.add_paragraph('conversation_id 通过 URL 查询参数传递，实现多轮对话上下文保持：')
    doc.add_paragraph('• 新用户：调用 create_conversation() 创建会话')
    doc.add_paragraph('• 老用户：从数据库获取 conversation_id 继续对话')
    doc.add_paragraph('• 会话 ID 存储在 MySQL users 表中')

    doc.add_heading('4.4 工作流变量', 2)
    doc.add_paragraph('通过 CozeVars 类定义传递给工作流的变量：')

    table_vars = doc.add_table(rows=5, cols=2)
    table_vars.style = 'Table Grid'

    var_header = table_vars.rows[0].cells
    var_header[0].text = '变量名'
    var_header[1].text = '说明'

    var_data = [
        ('buyer_name', '买家用户名'),
        ('product_title', '商品标题'),
        ('product_price', '商品价格'),
        ('order_status', '订单状态（简化后的状态值）'),
    ]

    for i, (var, desc) in enumerate(var_data, 1):
        row = table_vars.rows[i].cells
        row[0].text = var
        row[1].text = desc

    doc.add_paragraph()

    # ==================== 五、数据库模块 ====================
    doc.add_heading('五、数据库模块（对话记忆）', 1)

    doc.add_heading('5.1 功能描述', 2)
    doc.add_paragraph('使用 MySQL 存储用户信息和对话历史，实现跨会话的对话记忆功能。')

    doc.add_heading('5.2 数据表结构', 2)

    doc.add_paragraph('users 表：')
    table_users = doc.add_table(rows=5, cols=3)
    table_users.style = 'Table Grid'

    users_header = table_users.rows[0].cells
    users_header[0].text = '字段'
    users_header[1].text = '类型'
    users_header[2].text = '说明'

    users_data = [
        ('id', 'INT AUTO_INCREMENT', '主键'),
        ('buyer_name', 'VARCHAR(255)', '买家用户名（唯一）'),
        ('coze_conversation_id', 'VARCHAR(255)', 'Coze 会话 ID'),
        ('created_at / updated_at', 'DATETIME', '创建/更新时间'),
    ]

    for i, (field, type_, desc) in enumerate(users_data, 1):
        row = table_users.rows[i].cells
        row[0].text = field
        row[1].text = type_
        row[2].text = desc

    doc.add_paragraph()
    doc.add_paragraph('conversation_history 表：')

    table_history = doc.add_table(rows=6, cols=3)
    table_history.style = 'Table Grid'

    history_header = table_history.rows[0].cells
    history_header[0].text = '字段'
    history_header[1].text = '类型'
    history_header[2].text = '说明'

    history_data = [
        ('id', 'INT AUTO_INCREMENT', '主键'),
        ('buyer_name', 'VARCHAR(255)', '买家用户名'),
        ('role', 'VARCHAR(50)', '角色（user/assistant）'),
        ('content', 'TEXT', '消息内容'),
        ('coze_conversation_id', 'VARCHAR(255)', 'Coze 会话 ID'),
    ]

    for i, (field, type_, desc) in enumerate(history_data, 1):
        row = table_history.rows[i].cells
        row[0].text = field
        row[1].text = type_
        row[2].text = desc

    doc.add_paragraph()

    doc.add_heading('5.3 核心方法', 2)

    table_db = doc.add_table(rows=7, cols=2)
    table_db.style = 'Table Grid'

    db_header = table_db.rows[0].cells
    db_header[0].text = '方法'
    db_header[1].text = '功能'

    db_data = [
        ('connect()', '连接 MySQL 数据库'),
        ('init_tables()', '初始化数据表'),
        ('get_or_create_user()', '获取或创建用户'),
        ('update_conversation_id()', '更新用户的 Coze 会话 ID'),
        ('add_message()', '添加对话消息'),
        ('get_conversation_history()', '获取对话历史'),
    ]

    for i, (method, desc) in enumerate(db_data, 1):
        row = table_db.rows[i].cells
        row[0].text = method
        row[1].text = desc

    doc.add_paragraph()

    # ==================== 六、订单状态传递 ====================
    doc.add_heading('六、订单状态字段传递', 1)

    doc.add_heading('6.1 功能描述', 2)
    doc.add_paragraph('从闲鱼页面提取订单状态信息，作为辅助参数传递给 Coze 工作流，帮助 AI 更准确地理解用户意图和上下文。')

    doc.add_heading('6.2 状态提取来源', 2)
    doc.add_paragraph('订单状态从两个位置提取（优先级从高到低）：')
    doc.add_paragraph('1. 商品卡片（聊天区域右侧）- 优先使用')
    doc.add_paragraph('2. 会话列表（左侧列表项）- 备选')

    doc.add_heading('6.3 状态值映射表', 2)
    doc.add_paragraph('为避免传递动作指示性文字，将闲鱼原始状态映射为简化状态，同时记录系统消息内容：')

    # 状态映射表格
    table1 = doc.add_table(rows=14, cols=4)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    header_cells = table1.rows[0].cells
    header_cells[0].text = '闲鱼原始显示'
    header_cells[1].text = '传给 Coze 的值'
    header_cells[2].text = '系统消息内容'
    header_cells[3].text = '说明'

    # 数据行
    status_data = [
        ('等待买家付款', '待付款', '我已拍下，待付款', '订单已创建，买家未付款'),
        ('待付款', '待付款', '', '订单已创建，买家未付款'),
        ('已付款', '已付款', '', '买家已完成付款'),
        ('等待卖家发货', '已付款', '我已付款，等待你发货', '买家已付款，等待发货'),
        ('已发货', '已发货', '你已发货', '商品已寄出'),
        ('等待买家收货', '已发货', '已发货，等待买家确认', '商品在途中'),
        ('已收货', '已收货', '买家已确认收货', '买家已签收'),
        ('交易成功', '已完成', '交易成功', '交易结束'),
        ('交易关闭', '已关闭', '交易关闭', '交易关闭'),
        ('交易取消', '已取消', '订单已取消', '交易取消'),
        ('等待见面交易', '已付款', '', '线下交易（见面交易视为已付款）'),
        ('退款中', '退款中', '申请退款', '退款处理中'),
        ('已退款', '已退款', '退款成功', '退款完成'),
    ]

    for i, (orig, mapped, sys_msg, desc) in enumerate(status_data, 1):
        row = table1.rows[i].cells
        row[0].text = orig
        row[1].text = mapped
        row[2].text = sys_msg
        row[3].text = desc

    doc.add_paragraph()

    # ==================== 七、系统消息过滤 ====================
    doc.add_heading('七、系统消息过滤', 1)

    doc.add_heading('7.1 功能描述', 2)
    doc.add_paragraph('识别并过滤闲鱼系统通知消息（如下单、付款、发货通知），避免这些消息触发 AI 自动回复。')

    doc.add_heading('7.2 系统消息关键词', 2)
    doc.add_paragraph('以下消息内容会被标记为系统消息：')

    keywords = [
        '我已拍下，待付款',
        '我已付款，等待你发货',
        '请双方沟通及时确认价格',
        '请包装好商品',
        '你已发货',
        '已发货，等待买家确认',
        '买家已确认收货',
        '交易成功',
        '交易关闭',
        '订单已取消',
        '退款成功',
        '申请退款',
        '你撤回了一条消息',
        '对方撤回了一条消息',
        '对方正在输入',
    ]

    for kw in keywords:
        doc.add_paragraph(f'• {kw}')

    doc.add_paragraph()

    doc.add_heading('7.3 实现方式', 2)
    doc.add_paragraph('1. 消息数据类添加 is_system 字段')
    doc.add_paragraph('2. JavaScript 提取消息时检查关键词')
    doc.add_paragraph('3. 消息处理时跳过 is_system=True 的消息')

    doc.add_paragraph()

    # ==================== 八、日志系统 ====================
    doc.add_heading('八、日志系统', 1)

    doc.add_heading('8.1 日志配置', 2)
    doc.add_paragraph('使用 Loguru 框架，支持以下输出：')
    doc.add_paragraph('• 控制台输出：彩色格式化')
    doc.add_paragraph('• 文件输出：app_{date}.log（每日轮转，保留30天）')
    doc.add_paragraph('• 对话记录：conversations_{date}.log（保留90天）')

    doc.add_heading('8.2 日志级别优化', 2)

    # 日志优化表格
    table3 = doc.add_table(rows=5, cols=4)
    table3.style = 'Table Grid'

    log_header = table3.rows[0].cells
    log_header[0].text = '日志内容'
    log_header[1].text = '原级别'
    log_header[2].text = '新级别'
    log_header[3].text = '说明'

    log_data = [
        ('找到 X 个会话', 'INFO', 'DEBUG', '只在详细模式显示'),
        ('找到 X 个未读会话', 'INFO', '条件 INFO', '仅当 X > 0 时显示'),
        ('切换会话失败', 'ERROR', 'DEBUG', '降级为调试日志'),
        ('获取商品信息失败', 'ERROR', 'DEBUG', '降级为调试日志'),
    ]

    for i, (content, old, new, desc) in enumerate(log_data, 1):
        row = table3.rows[i].cells
        row[0].text = content
        row[1].text = old
        row[2].text = new
        row[3].text = desc

    doc.add_paragraph()

    # ==================== 九、GUI 界面 ====================
    doc.add_heading('九、GUI 界面', 1)

    doc.add_heading('9.1 界面功能', 2)
    doc.add_paragraph('• 配置管理：API Token、Bot ID、检查间隔')
    doc.add_paragraph('• 数据库配置：MySQL 连接参数')
    doc.add_paragraph('• 启动/停止控制')
    doc.add_paragraph('• 实时日志显示')
    doc.add_paragraph('• 详细日志开关')

    doc.add_heading('9.2 窗口配置', 2)
    doc.add_paragraph('默认窗口尺寸: 600x800')
    doc.add_paragraph('最小尺寸: 520x700')
    doc.add_paragraph('允许用户自由拉伸窗口大小')

    doc.add_heading('9.3 详细日志开关', 2)
    doc.add_paragraph('添加"显示详细日志"复选框，切换日志级别：')
    doc.add_paragraph('• 关闭（默认）: 只显示 INFO 及以上级别')
    doc.add_paragraph('• 开启: 显示 DEBUG 及以上级别')

    doc.add_paragraph()

    # ==================== 十、配置说明 ====================
    doc.add_heading('十、配置说明', 1)

    doc.add_heading('10.1 环境变量', 2)

    table_env = doc.add_table(rows=10, cols=3)
    table_env.style = 'Table Grid'

    env_header = table_env.rows[0].cells
    env_header[0].text = '变量名'
    env_header[1].text = '默认值'
    env_header[2].text = '说明'

    env_data = [
        ('COZE_API_TOKEN', '（必填）', 'Coze API 访问令牌'),
        ('COZE_BOT_ID', '（必填）', 'Coze Bot ID'),
        ('XIANYU_CHECK_INTERVAL', '10', '消息检查间隔（秒）'),
        ('HEADLESS', 'false', '是否无头模式运行浏览器'),
        ('DB_HOST', 'localhost', 'MySQL 数据库地址'),
        ('DB_PORT', '3306', 'MySQL 端口'),
        ('DB_USER', 'root', '数据库用户名'),
        ('DB_PASSWORD', 'root', '数据库密码'),
        ('DB_NAME', 'xianyu', '数据库名称'),
    ]

    for i, (var, default, desc) in enumerate(env_data, 1):
        row = table_env.rows[i].cells
        row[0].text = var
        row[1].text = default
        row[2].text = desc

    doc.add_paragraph()

    # ==================== 十一、Coze 工作流配置建议 ====================
    doc.add_heading('十一、Coze 工作流配置建议', 1)
    doc.add_paragraph('为确保 AI 正确使用订单状态信息，建议在 Coze 工作流提示词中添加以下规则：')
    doc.add_paragraph()
    doc.add_paragraph('【重要规则】')
    doc.add_paragraph('1. 首先理解并直接回答用户的问题')
    doc.add_paragraph('2. 只有当用户询问订单/物流/付款相关问题时，才提及订单状态')
    doc.add_paragraph('3. 对于简单问候（如"在吗"、"你好"），直接友好回应即可')
    doc.add_paragraph()
    doc.add_paragraph('【可用信息】')
    doc.add_paragraph('• 买家名称: {{buyer_name}}')
    doc.add_paragraph('• 商品标题: {{product_title}}')
    doc.add_paragraph('• 商品价格: {{product_price}}')
    doc.add_paragraph('• 订单状态: {{order_status}}')
    doc.add_paragraph()
    doc.add_paragraph('【回复示例】')
    doc.add_paragraph('用户: 在吗')
    doc.add_paragraph('回复: 在的亲，有什么可以帮您的吗？')
    doc.add_paragraph()
    doc.add_paragraph('用户: 发货了吗')
    doc.add_paragraph('回复: 亲，您的订单当前状态是{{order_status}}，请耐心等待哦~')

    doc.add_paragraph()

    # ==================== 十二、文件修改清单 ====================
    doc.add_heading('十二、文件清单', 1)

    table4 = doc.add_table(rows=9, cols=2)
    table4.style = 'Table Grid'

    file_header = table4.rows[0].cells
    file_header[0].text = '文件'
    file_header[1].text = '主要功能'

    file_data = [
        ('gui.py', 'GUI 界面、配置管理、启动控制、日志显示'),
        ('xianyu_browser.py', '浏览器自动化、页面操作、数据提取、订单状态映射'),
        ('message_handler.py', '消息处理核心、系统消息过滤、订单状态传递'),
        ('coze_client.py', 'Coze API 客户端、会话管理、变量传递'),
        ('db_manager.py', 'MySQL 数据库管理、用户存储、对话历史'),
        ('config.py', '配置管理、CozeVars 变量定义'),
        ('logger_setup.py', '日志配置、对话记录'),
        ('main.py', '命令行入口'),
    ]

    for i, (file, content) in enumerate(file_data, 1):
        row = table4.rows[i].cells
        row[0].text = file
        row[1].text = content

    doc.add_paragraph()

    # ==================== 十三、版本历史 ====================
    doc.add_heading('十三、版本历史', 1)

    table5 = doc.add_table(rows=3, cols=3)
    table5.style = 'Table Grid'

    ver_header = table5.rows[0].cells
    ver_header[0].text = '版本'
    ver_header[1].text = '日期'
    ver_header[2].text = '更新内容'

    ver_data = [
        ('v1.0', '-', '初始版本：浏览器自动化、Coze AI 集成、GUI 界面、MySQL 对话记忆'),
        ('v1.1', '2026-01-06', '订单状态传递、系统消息过滤、日志优化、GUI 改进（窗口可拉伸、详细日志开关）'),
    ]

    for i, (ver, date, content) in enumerate(ver_data, 1):
        row = table5.rows[i].cells
        row[0].text = ver
        row[1].text = date
        row[2].text = content

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('— 文档结束 —')

    # 保存文档
    output_path = r'C:\Users\pc\Desktop\RPA\docs\闲鱼智能客服RPA_技术方案文档.docx'
    doc.save(output_path)
    print(f'文档已保存: {output_path}')
    return output_path

if __name__ == '__main__':
    create_document()
